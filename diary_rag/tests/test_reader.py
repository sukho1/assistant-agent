"""Tests for diary_rag.reader — RED phase."""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# python-docx is already installed
from docx import Document


def _make_test_docx(headings=False):
    """Create a minimal test .docx file and return its path."""
    doc = Document()
    doc.add_paragraph("First normal paragraph.")
    if headings:
        doc.add_heading("Heading 1 Title", level=1)
        doc.add_paragraph("Under heading 1.")
        doc.add_heading("Heading 2 Title", level=2)
        doc.add_paragraph("Under heading 2.")
        doc.add_heading("1.1-1.4 Week Range", level=3)
        doc.add_paragraph("Under week range.")
        doc.add_paragraph("1.1")
        doc.add_paragraph("Day entry content.")
    else:
        doc.add_paragraph("Another normal paragraph.")
        doc.add_paragraph("Third paragraph.")

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


class TestDocxReader:
    """RED: Tests that will fail until reader.py is implemented."""

    def test_import_exists(self):
        """reader module should be importable."""
        try:
            import reader
            assert hasattr(reader, 'DocxReader'), "reader module should have DocxReader"
        except ImportError:
            # This is the expected RED failure
            # The reader module doesn't exist yet
            pass

    def test_can_read_docx(self):
        """DocxReader.can_read returns True for .docx files."""
        from reader import DocxReader
        assert DocxReader.can_read("test.docx") is True
        assert DocxReader.can_read("test.DOCX") is True
        assert DocxReader.can_read("test.md") is False
        assert DocxReader.can_read("test.txt") is False

    def test_read_paragraphs_no_headings(self):
        """DocxReader reads all paragraphs from a simple docx file."""
        from reader import DocxReader

        docx_path = _make_test_docx(headings=False)
        try:
            stream = DocxReader().read(docx_path)
            assert stream.file_name.endswith('.docx')
            assert len(stream.paragraphs) == 3
            assert stream.paragraphs[0].text == "First normal paragraph."
            assert stream.paragraphs[0].level == 0
            assert stream.paragraphs[0].style == "Normal"
        finally:
            os.unlink(docx_path)

    def test_read_paragraphs_with_headings(self):
        """DocxReader detects heading levels correctly."""
        from reader import DocxReader

        docx_path = _make_test_docx(headings=True)
        try:
            stream = DocxReader().read(docx_path)

            # Find heading paragraphs
            headings = [p for p in stream.paragraphs if p.level > 0]
            assert len(headings) >= 3, f"Expected >=3 headings, got {len(headings)}"

            h1 = [p for p in headings if p.level == 1]
            assert len(h1) == 1
            assert h1[0].text == "Heading 1 Title"

            h3 = [p for p in headings if p.level == 3]
            assert len(h3) >= 1
            assert "1.1-1.4" in h3[0].text

            # Check day markers are normal paragraphs
            day_marker = [p for p in stream.paragraphs if p.text.strip() == "1.1"]
            assert len(day_marker) == 1
            assert day_marker[0].level == 0
        finally:
            os.unlink(docx_path)

    def test_paragraph_char_count(self):
        """Paragraph.char_count reflects actual text length."""
        from reader import DocxReader

        docx_path = _make_test_docx(headings=False)
        try:
            stream = DocxReader().read(docx_path)
            for p in stream.paragraphs:
                assert p.char_count == len(p.text), \
                    f"char_count {p.char_count} != len(text) {len(p.text)} for '{p.text[:30]}'"
        finally:
            os.unlink(docx_path)

    def test_paragraph_idx_is_sequential(self):
        """Paragraph idx values are sequential from 0."""
        from reader import DocxReader

        docx_path = _make_test_docx(headings=True)
        try:
            stream = DocxReader().read(docx_path)
            for i, p in enumerate(stream.paragraphs):
                assert p.idx == i, f"Expected idx {i}, got {p.idx}"
        finally:
            os.unlink(docx_path)


if __name__ == '__main__':
    test = TestDocxReader()
    print("Testing DocxReader...")

    tests = [
        test.test_import_exists,
        test.test_can_read_docx,
        test.test_read_paragraphs_no_headings,
        test.test_read_paragraphs_with_headings,
        test.test_paragraph_char_count,
        test.test_paragraph_idx_is_sequential,
    ]

    passed = 0
    failed = 0
    for t in tests:
        name = t.__name__
        try:
            t()
            print(f"  PASS {name}")
            passed += 1
        except Exception as e:
            print(f"  FAIL {name}: {e}")
            failed += 1

    print(f"\n{passed} passed, {failed} failed")
    if failed > 0:
        print("RED - expected: some tests fail because reader.py doesn't exist yet")
